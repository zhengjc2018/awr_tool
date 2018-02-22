# -*- coding: utf-8 -*-  
import os,docx,bs4,re,argparse,datetime,time,string,functools,sys
from bs4 import BeautifulSoup
from time import sleep
from docx import Document
from docx.shared import RGBColor
from docx.shared import Pt
import matplotlib.pyplot as plt
import matplotlib.pylab as pylab
from functools import reduce
from init import init_para_cmplist
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.dml import MSO_THEME_COLOR
import unittest
from docx2pdf import Doc2PDF

################################################ global variable ################################################################################
# 设置全局的文字大小/粗细/颜色/斜体
# global document
document=Document()
# style=document.styles['Normal']
style1=document.styles['Normal']
paragraph= document.add_paragraph()
paragraph_format=paragraph.paragraph_format
font=style1.font
font.name="Calibri"
font.size = Pt(8)

paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
paragraph.space_after = Pt(3)
paragraph.space_before = Pt(3)
################################################################################################################################################
baseloca=os.path.dirname(os.getcwd())  #app同级目录
tmploca=os.path.join(baseloca,'templates')
resloca=os.path.join(baseloca,'result')
global file
rac = "单机" 
advise_flag=0
################################################################################################################################################
################################################################################################################################################
################################################################################################################################################
#########################                             辅助函数部分                                                     ######################### 
################################################################################################################################################
################################################################################################################################################
def datechange(*args):
	date_dict = {'Jan':'01','Feb':'02','Mar':'03',"Apr":'04','May':'05','Jun':'06',\
				'Jul':'07','Aug':'08','Sep':'09','Oct':'10','Nov':'11','Dec':'12'}
	tmp1=args[0].split(' ')
	t1 = list(reversed(tmp1[0].split('-')))
	t2 = tmp1[1].split(':')[:2]
	t3 = args[2] + '数据库性能分析报告(20'+t1[0] + date_dict[t1[1]] + t1[2] + ''.join(t2) + '-' +''.join(args[1].split(' ')[1].split(':')[:2]) +').docx'
	return t3

# args[1]传入列的个数,args[2]传入th标签的内容,args[3]传入td标签内容,args[0]为document实例的，作用：生成表格
def generate_table(*args,autofit=0,line=0):
	if args[3]:
		table = args[0].add_table(rows=1, cols=args[1],style='Light Shading Accent 1')
		table._tblPr.autofit=True
		table.autofit=True
		######################## 设置表格内单元格的边距，按每列字符串最大长度a*(100/每列最大长度相加之和)*100000  #####################################
		tmp1=args[3].copy()
		tmp1.append(args[2])
		# tmp1=args[3]+args[2]
		tmp=[]
		for i in range(len(args[2])):
			tmp.append( [len(str(j[i])) for j in tmp1])
		tmp2=[float(max(i)) for i in tmp]
		total_length=reduce(lambda x,y:x+y,tmp2)
		row_size=[i*(10000000/total_length) for i in tmp2]   #获得每列的行间距

		hdr_cells = table.rows[0].cells  #获得第1个表格
		for i,j in enumerate(args[2]): 	# 设置标题行的内容
			if j==None:
				j=' '
			else:
				pass
			hdr_cells[i].text = str(j)
			if autofit:
				hdr_cells[i].width = row_size[i]
			else:
				pass
		
		for k in args[3]: 	# 设置数据行内容
			row_cells = table.add_row().cells
			for i,j in enumerate(k):
				row_cells[i].text=str(j)
		if line:
			document.add_paragraph(" ")
		global advise_flag
		advise_flag=1
	else:
		pass

# 判断是否为不正常，假如是加红，然后写到最开始的表格中，第4个参数：1为小于参考值正常，0为大于参考值正常
def functionB(*args):
	row,col=args[2]
	try:
		if args[3]:
			if float(args[0]) > args[1] :
				run=document.tables[0].cell(row,col).paragraphs[0].add_run('不正常')
				run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
			else:
				run=document.tables[0].cell(row,col).paragraphs[0].add_run('正常')
		else:
			if float(args[0]) < args[1] :
				run=document.tables[0].cell(row,col).paragraphs[0].add_run('不正常')
				run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
			else:
				run=document.tables[0].cell(row,col).paragraphs[0].add_run('正常')		
	except:
		run=document.tables[0].cell(row,col).paragraphs[0].add_run('N/A')
		run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

# 返回一个table里面的所有内容
# @functools.lru_cache(maxsize=16)
def functionC(*args,key_word=[],zipnum=[],th_len=0):
	summary=args[1]
	soup=args[0].find_all("table",summary=re.compile(summary))
	try:
		tmp_td = [i.string.replace(',','').strip() for i in soup[0].find_all("td")]
		len_td = len(tmp_td)
		total_table_th = [i.string for i in soup[0].find_all("th")]
		len_th =th_len if th_len else  len(total_table_th)
		if zipnum:
			total_table_td = []
			for num in zipnum:
				total_table_td.append([j for i,j in enumerate(tmp_td) if i%len_th == num ])
			return total_table_th,total_table_td
		else:
			th_length = len(total_table_th) if len(total_table_th) else th_len
			total_table_td = [ tmp_td[i:i+len_th] for i in range(0,len_td,th_length)]
		# 寻找关键字对应的值
			value=[]
			if key_word:
				tmp=dict(zip(total_table_th,total_table_td[0]))
				for i in key_word:
					value.append(tmp.get(i,'N/A'))
				return total_table_th,total_table_td,value
			return total_table_th,total_table_td
	except:
		return [],[],key_word

# 获得的buffer pool中的值可能含有keep pool的值，所以进行筛选
def functionA(*args):
	flag=0
	l=args[0]
	for i,j in enumerate(l,1):
		try:
			if l[flag] < l[i]:
				flag+=1
			else:
				return l[:i]
		except:
			pass
	return l

# 画图，并将图片插入到docx文件中
def GetGraph(*args):
	table_summary=args[1]
	keyword=args[2]
	picname=os.path.basename(file).replace('docx','png')
	picloca=os.path.join(tmploca,picname)
	document.add_paragraph(args[3],style="Heading 9")
	try:
		soup=args[0].find_all("table",summary=re.compile(table_summary))
		aaa=soup[1] if len(soup)>1 else soup[0]
		t=[i.string.replace(' ','') for i in aaa.find_all('th')]
		total_len=len(t)
		start_len=t.index(keyword[0].replace(' ',''))
		end_len=t.index(keyword[1].replace(' ',''))
		tmp=aaa.find_all('td')
		x1=[float(j.string.replace(',','')) for i,j in enumerate(tmp) if i%total_len == start_len]
		y1=[float(j.string.replace(',','')) for i,j in enumerate(tmp) if i%total_len == end_len]
		tx1=functionA(x1)
		ty1=y1[:len(tx1)]
		################################### 绘制折线图 #################################################
		plt.close()  		#清图用，否则会和上一次的一起显示
		sleep(0.5)
		plt.figure(figsize=(6,4))  
		plt.plot(tx1,ty1,label='%s'%args[3],linewidth=3,color='r',marker='o',markerfacecolor='blue',markersize=4) 
		plt.xlabel(keyword[0]) 
		plt.ylabel(keyword[1]) 
		plt.title(args[3]) 
		# plt.grid(True)	#添加网格线
		plt.legend()    	#增加图例 
		# plt.subplots_adjust(left=0.19,right=1,wspace=0.25,hspace=0.25,bottom=0.13,top=0.91)  
		plt.subplots_adjust(left=0.15,right=0.99,top=0.92)  
		plt.savefig(picloca,dpi=100)
		# document.add_picture(picloca,width=Inches(1.25))
		document.add_picture(picloca)
		document.add_paragraph("当前%s的配置为:%sMB"%(args[4][0],args[4][1]),style="List Bullet 2")	
		document.save(file)
		os.remove(picloca)
	except:
		document.add_paragraph("No Data found",style="List Bullet 2")

# 将单位和数值分离，都转化为ms，float类型
def CovertUsToMs(argv):
	if 'ms' in argv:
		tmp=argv.split('ms')
		return (round(float(argv.split('ms')[0].replace(',','').strip()),4))
	elif 'us' in argv:
		return (round(float(argv.split('us')[0].replace(',','').strip())/1000,4))
	else:
		return round(float(argv),4)

################################################################################################################################################
################################################################################################################################################
#########################                             主功能函数部分												   ######################### 
################################################################################################################################################
################################################################################################################################################
# 第一部分：AWR报告解析总结，显示mainreport中的内容
def MainReport():
	title=["AWR报告解析总结","AWR报告概况","主机资源概况","数据库内存配置","会话登录阶段","SQL解析阶段","SQL执行阶段","事务提交阶段",\
			"RAC Statistics","数据库参数建议"]
	document.add_heading("Main Report")
	for i in title:
		document.add_paragraph(i,style='List Bullet 2')

	document.add_heading("一.AWR报告解析总结")
	title=['序号','检查流程','细项','检查结果']
	t1=['1','主机资源','CPU资源','']
	t2=[' ',' ','I/O资源','']
	t3=[' ',' ','内存资源','']
	t4=[' ',' ','网络资源','']
	t5=['2','数据库内存配置','Buffer cache','']
	t6=[' ',' ','Shared pool','']
	t7=[' ',' ','PGA','']
	t8=['3','会话登录阶段','会话连接时间','']
	t9=[' ',' ','登录次数','']
	t10=['4','SQL解析阶段','解析','']
	t11=[' ',' ','硬解析','']
	t12=['5','SQL执行阶段','执行时间','']
	t13=[' ',' ','逻辑读','']
	t14=[' ',' ','物理读','']
	t15=['6','事务提交阶段','提交响应时间','']
	t16=['7','RAC统计','集群响应时间','']
	t17=[' ',' ','节点内部通信性能','']
	t18=[' ',' ','节点内部通信PING延迟','']
	t19=['8','数据库参数建议','是否存在建议修改参数','']
	
	content=[t1,t2,t3,t4,t5,t6,t7,t8,t9,t10,t11,t12,t13,t14,t15,t16,t17,t18,t19]
	generate_table(document,4,title,content)
	# 设置行间距以及字体大小
	paragraph=document.add_paragraph()
	paragraph_format = paragraph.paragraph_format
	paragraph_format.space_before = Pt(5)
	paragraph_format.space_after = Pt(5)
	run=paragraph.add_run("注：检查结果为不正常的细项可以在对应的流程中进一步查看详情。")
	run.font.color.theme_color = MSO_THEME_COLOR.ACCENT_1
	run.font.size = Pt(7)
	# run.italic = True

	# 合并单元格
	# document.tables[0].cell(1,0).merge(document.tables[0].cell(4,0))
	# document.tables[0].cell(1,1).merge(document.tables[0].cell(4,1))
	# document.tables[0].cell(6,0).merge(document.tables[0].cell(7,0))

#第二部分：包括DB Name开始到Cursors/Session之间的内容,*title为每个表的标题栏，*result为表数据
def BasicSituation(*args):
	document.add_heading("二.AWR报告概况")
	DBnametitle,DBnameResult,dbvalue=functionC(args[0],"database instance information",key_word=['DB Name','Release','RAC'])
	HostNametitle,HostNameResult,hostvalue=functionC(args[0],"host information",key_word=["Platform","CPUs","Memory (GB)"])
	Snaptitle,SnapResult=functionC(args[0],"snapshot information")

	db_name,release,YesOrNo = dbvalue
	global rac
	rac='RAC' if (YesOrNo == 'YES') else '单机'
	platform,cpus,memory = hostvalue

	tmp_save_file=datechange(SnapResult[0][2],SnapResult[1][2],db_name)
	global file
	file = os.path.join(resloca,tmp_save_file)

	generate_table(document,len(DBnametitle),DBnametitle,DBnameResult,autofit=1,line=1)
	generate_table(document,len(HostNametitle),HostNametitle,HostNameResult,autofit=1,line=1)
	generate_table(document,len(Snaptitle),Snaptitle,SnapResult,autofit=1)
    #####################################  输出说明  #############################################################################3
	document.add_paragraph("输出说明:",style="Heading 9")
	document.add_paragraph("当前数据库名字为:%s，数据库版本为%s，为%s架构。"%\
							(db_name,release,rac),style="List Bullet 2")
	document.add_paragraph("当前数据库操作系统为:%s，有%s颗逻辑CPU，配置了%sGB内存。"%\
							(platform,cpus,memory),style="List Bullet 2")
	document.add_paragraph("数据库性能采样开始时间为:%s，结束时间为:%s，采样间隔为%s。"%\
							(SnapResult[0][2],SnapResult[1][2],SnapResult[2][2]),style="List Bullet 2")
	try:
		d_value=abs(float(SnapResult[1][3])-float(SnapResult[0][3]))
		sug="会话连接数整体波动不大，总体表现正常" if d_value<100 else "会话连接数量波动较大"
		document.add_paragraph("采样开始时,数据库连接数为:%s个，采样结束时为%s个，两者相差%s个，%s。"%\
									(SnapResult[0][3],SnapResult[1][3],d_value,sug),style="List Bullet 2")
	except:
		pass
	try:   		# 消耗的逻辑cpu占总cpu超过75%，则认为紧张
		time=SnapResult[3][2]
		log_cpus=round((float(time.replace('(mins)',''))/float(SnapResult[2][2].replace('(mins)',''))),3)
		sug = "当前CPU资源比较紧张。" if log_cpus/float(cpus.replace(',','')) > 0.75 else "CPU资源使用正常。"
		document.add_paragraph("数据库性能采样期间，DB Time消耗%s，消耗了%s颗逻辑CPU，%s"%(time,log_cpus,sug),style="List Bullet 2")
		# print(time,SnapResult[2][2],cpus,log_cpus)
	except:
		pass

#第三部分：主机资源
def HostResource(*args):
	document.add_heading("三.主机资源概况")
	# 1.获取平均使用时间/平均空闲时间/平均空闲率 AVG_BUSY_TIME AVG_IDLE_TIME AVG_FREE
	try:
		*a,tmp = functionC(args[0],"operating systems statistics",zipnum=[0,1])
		Statistic = tmp[0]
		Value  = [float(i.replace(',','')) for i in tmp[1]]
		OperateSysStat= dict(zip(Statistic,Value))
		SYS_TIME=OperateSysStat.get('SYS_TIME')/60
		USER_TIME=OperateSysStat.get('USER_TIME')/60
		AVG_BUSY_TIME = round( (SYS_TIME + USER_TIME)/4,2)
		AVG_IDLE_TIME = round(OperateSysStat.get('IDLE_TIME')/240,2)
		AVG_FREE = (round((AVG_IDLE_TIME)/(AVG_BUSY_TIME+AVG_IDLE_TIME)*100,3)) if (AVG_BUSY_TIME is not 'N/A' and AVG_IDLE_TIME is not 'N/A') else 'N/A'
	except:
		AVG_BUSY_TIME = AVG_IDLE_TIME=AVG_FREE='N/A'

	# 2.获取I/O资源信息，包括平均iops/平均每秒吞吐量/最大相应时间/平均响应时间  Total Requests=Total_req,Total (MB)=Total_MB
	try:
		*a,tmp = functionC(args[0],"This table displays Instance activity statistics",th_len=3)
		Total_req = Total_MB = 0
		for i in tmp:
			if i[0] in ["physical read IO requests","physical write IO requests"]:
				Total_req += float(i[2].replace(",",""))
			elif i[0] in ["physical read bytes","physical write bytes"]:
				Total_MB += float(i[2].replace(",",""))/1048576
			else:
				pass

		*a,tmp = functionC(args[0],"This table displays non-key Instance activity statistics",th_len=3)
		for i in tmp:
			if i[0] in ["physical read IO requests","physical write IO requests"]:
				Total_req += float(i[2].replace(",",""))
			elif i[0] in ["physical read bytes","physical write bytes"]:
				Total_MB += float(i[2].replace(",",""))/1048576
			else:
				pass
	except:
		Total_req=Total_MB='N/A'

	# 获得平均响应时间和最大响应时间
	try:
		tmp01=tmp02=0
		t=[]
		*_,tmp=functionC(args[0],"foreground wait class statistics")
		for i in tmp:
			if i[0] in ["User I/O","System I/O"]:
				tmp_value = float(i[1].replace(',',''))
				tmp_time  = 1000 * float(i[3].replace(',',''))
				tmp01 += tmp_value
				tmp02 += tmp_time
				try:
					t.append(round(tmp_time/tmp_value,3))
				except:
					t.append(0)
		AvRd_MS=max(t)
		Avg_time=round(tmp02/tmp01,3)
	except:
		AvRd_MS=Avg_time='N/A'
	
	#获得内存使用率 Menstatis
	try:
		*a,tmp = functionC(args[0],"This table displays memory statistics")
		Menstatis = [max(float(i[1]),float(i[2])) for i in tmp if "% Host Mem used for SGA+PGA:" in i][0]
	except:
		Menstatis = 'N/A'

	# 获得每秒私网流量CacheLoadProfile
	try:
		*a,tmp=functionC(args[0],"information about global cache load",zipnum=[0,1])
		Estd_traffic = [j.strip() for j in tmp[0]]
		PerSedvalue  = [j.replace(',','').replace('\xa0','0').strip() for j in tmp[1]]
		l=dict(zip(Estd_traffic,PerSedvalue))
		CacheLoadProfile = l.get('Estd Interconnect traffic (KB)','N/A')
	except:
		CacheLoadProfile = 'N/A'
	#############################   判断是正常还是不正常   #############################
	functionB(AVG_FREE,30,(1,3),0)
	functionB(AvRd_MS,20,(2,3),1)
	functionB(Menstatis,90,(3,3),1)
	functionB(CacheLoadProfile,200480,(4,3),1)
	####################################### 表格 #######################################
	t0=['检查流程','细项','数值']
	t1=['CPU资源','平均使用时间(min)',AVG_BUSY_TIME]
	t2=['		','平均空闲时间(min)',AVG_IDLE_TIME]
	t3=['		','平均空闲率(%)',AVG_FREE]
	t4=['I/O资源','平均IOPS',round(Total_req,3)]
	t5=['		','平均每秒吞吐量(MB)',round(Total_MB,3)]
	t6=['		','最大响应时间(ms)',AvRd_MS]
	t7=['		','平均响应时间(ms)',Avg_time]
	t8=['内存资源','内存使用率(%)',Menstatis]
	t9=['网络资源','每秒私网流量(KB)',CacheLoadProfile]
	generate_table(document,3,t0,[t1,t2,t3,t4,t5,t6,t7,t8,t9])
	document.add_paragraph("建议:",style="Heading 9")
	####################################### 建议 #######################################
	try:
		tmp = "当前CPU平均空闲率为%s%%，"%(AVG_FREE)
		tmp += "空闲率低于30%，CPU资源紧张。"  if float(AVG_FREE)<30  else "空闲率高于30%，CPU资源正常。"
		document.add_paragraph(tmp,style="List Bullet 2")
	except:
		pass

	try:
		tmp = "当前IO最大响应时间为%sms/次，"%AvRd_MS
		if float(AvRd_MS)<20 :
			tmp += "低于参考值20.0ms/次，IO资源正常。"
		elif float(AvRd_MS)<40:
			tmp += "高于参考值20.0ms/次，IO资源可能存在异常。"
		else:
			tmp += "高于参考值40.0ms/次，IO资源存在异常。"
		document.add_paragraph(tmp,style="List Bullet 2")
	except:
		pass

	try:
		tmp="当前内存使用率为:%s%%，"%Menstatis
		tmp += "低于参考值80%，内存资源使用正常。" if float(Menstatis)<80  else "高于参考值80%，内存资源使用不正常。"
		document.add_paragraph(tmp,style="List Bullet 2")
	except:
		pass

	try:
		tmp="当前私网流量为:%sKB/s，"%CacheLoadProfile
		tmp += "高于参考值20MB/s，流量不正常。" if float(CacheLoadProfile)>20480  else "低于参考值20MB/s，流量正常。"
		document.add_paragraph(tmp,style="List Bullet 2")
	except:
		pass

# 第四部分：数据库内存配置
def MemoryConfig(*args):
	document.add_heading("四.数据库内存配置")
	################################  BUFFER POOL取值  ################################
	try:
		*a,tmp = functionC(args[0],"memory dynamic component statistics",zipnum=[0,2])
		comment=tmp[0]
		maxsize=[float(j.replace(',','')) for j in tmp[1]]
		l=dict(zip(comment,maxsize))
		DEFAULT_BUFFER_CACHE=l.get('DEFAULT buffer cache')
		KEEP_BUFFER_CACHE=l.get('KEEP buffer cache')
		RECYCLE_BUFFER_CACHE=l.get('RECYCLE buffer cache')
		SHARED_POOL=l.get('shared pool')
		SGA_TARGET=l.get('SGA Target')
		# print(DEFAULT_BUFFER_CACHE,KEEP_BUFFER_CACHE,RECYCLE_BUFFER_CACHE)
	except:
		RECYCLE_BUFFER_CACHE=DEFAULT_BUFFER_CACHE=KEEP_BUFFER_CACHE='N/A'
	################################  SGA取值  ################################
	try:
		*a,tmp = functionC(args[0],"name and value of init.ora parameters")
		try:
			SGA_MAX_VALUE = [ float(i[1]) for i in tmp if "sga_max_size" in i[0] ][0]
		except:
			pass
		sugg0="该数据库SGA内存管理方式为手动管理。" if SGA_TARGET == 0.0 else "该数据库SGA内存管理方式为自动管理。"
		SGA_VALUE = round(SGA_MAX_VALUE/1048576,2) if SGA_TARGET == 0.0 else SGA_TARGET
	except:
		SGA_VALUE='N/A'
		sugg0=''

	try:
		*_,tmp = functionC(args[0],"instance efficiency percentages",zipnum=[0,1],th_len=2)
		a=[j.replace(' ','') for j in tmp[0]]
		b=[j.strip() for j in tmp[1]]
		l=dict(zip(a,b))

		BUFFER_HIT=float(l.get('BufferHit%:'))
		LIBRARY_HIT=float(l.get('LibraryHit%:'))
		SP_HIT=float(l.get("SoftParse%:"))
	except:
		LIBRARY_HIT=BUFFER_HIT=SP_HIT='N/A'
	################################  建议  ################################
	sugg1="No Data found" if BUFFER_HIT == "N/A" else ("当前buffer cache命中率为:%s%%，"%BUFFER_HIT,\
														'低于参考值90%，不正常。' if (BUFFER_HIT)<90 else "高于参考值90%，正常。")
	sugg2="No Data found" if LIBRARY_HIT == "N/A" else ("当前Library Cache命中率为:%s%%，"%LIBRARY_HIT,\
														'低于参考值98%，不正常。' if (LIBRARY_HIT)<98 else "高于参考值98%，正常。")	
	sugg3="No Data found" if LIBRARY_HIT == "N/A" else ("当前keep pool命中率为:%s%%，"%BUFFER_HIT,\
														'低于参考值99%，不正常。' if (BUFFER_HIT)<99 else "高于参考值99%，正常。")

	functionB((BUFFER_HIT),90,(5,3),0)
	functionB((SP_HIT),98,(6,3),0)
	################################  PGA取值  ################################
	try:
		*_ ,tmp = functionC(args[0],"memory dynamic component statistics")
		PGA_USE=[float(i[2]) for i in tmp if "PGA Target" in i[0]] [0]
	except:
		PGA_USE='N/A'

	try:
		*_ , tmp = functionC(args[0],"shared pool advisory. Size factor, estimated library cache size",zipnum=[1,2])
		t1=dict(zip(tmp[0],tmp[1]))
		EST_LC_SIZE=t1.get('1.00')
	except:
		EST_LC_SIZE='N/A'

	################################  表格  ###################################
	document.add_heading('1.数据库内存资源',level=2)
	t1=['内存组件','细项','数值(MB)','命中率(%)']
	t0=['SGA','/',SGA_VALUE,'N/A']
	t2=['Buffer cache','Default Pool',DEFAULT_BUFFER_CACHE,BUFFER_HIT]
	t3=['		','Keep Pool',KEEP_BUFFER_CACHE,BUFFER_HIT]
	t4=['		','Recycle Pool',RECYCLE_BUFFER_CACHE,'N/A']
	t5=['Shared Pool','/',SHARED_POOL,'N/A']
	t6=['		','Library Cache',EST_LC_SIZE,LIBRARY_HIT]
	t7=['		','Dictionary Cache','N/A','N/A']
	t8=['PGA','/',PGA_USE,'N/A']
	generate_table(document,4,t1,[t0,t2,t3,t4,t5,t6,t7,t8],autofit=1)
	document.add_paragraph("建议:",style="Heading 9")
	if sugg0:
		document.add_paragraph(sugg0,style="List Bullet 2")
	else:
		pass
	document.add_paragraph(sugg1,style="List Bullet 2")
	document.add_paragraph(sugg2,style="List Bullet 2")
	document.add_paragraph(sugg3,style="List Bullet 2")

	###################################  内存抖动信息  ##########################################
	document.add_heading('2.SGA内存抖动信息',level=2)
	try:
		*_ , tmp1 = functionC(args[0],"memory dynamic component statistics")
		t1=['内存组件','开始大小(MB)','当前大小(MB)','最小值(MB)','最大值(MB)','操作次数','操作类型']
		generate_table(document,7,t1,tmp1,autofit=1)
		comp_list=[i[0] for i in tmp1]
		opercount_list=[float(i[5]) for i in tmp1]
		maxsize_list=[float(i[4].replace(',','')) for i in tmp1]
		document.add_paragraph("建议:",style="Heading 9")
		tmp=[]
		try:
			for i,j in enumerate(opercount_list):
				if not(j == 0):
					tmp.append(maxsize_list[i])
				else:
					pass
			if tmp:
				document.add_paragraph("SGA内存配置存在抖动现象,内存抖动对于高响应需求的OLTP业务影响较大，需要避免，可以根据内存配置曲线或咨询DBA进行调整。",style="List Bullet 2")
			else:
				document.add_paragraph("SGA组件之间没有发生抖动，表现正常",style="List Bullet 2")
		except:
			pass
	except:
		document.add_paragraph("No Data found",style="List Bullet 2")
	###################################  PGA溢出信息  ##########################################
	document.add_heading('3.PGA溢出信息',level=2)
	try:
		flag=0
		a,value = functionC(args[0],"PGA aggregate target histograms")
		title=[i.strip() for i in a]
		t=[i[-2:] for i in value]
		generate_table(document,len(title),title,value)
		for j in t:
			if j.count('0') < 2:
				flag+=1
			else:
				pass
		document.add_paragraph("建议:",style="Heading 9")
		if flag:
			document.add_paragraph('在1G和2G内存区间内，总的1-Pass Execs或M-Pass Execs出现的次数为:%s次。当前PGA配置为:%sMB，PGA内存配置可能不足。'\
									%(flag,PGA_USE),style="List Bullet 2")
			run=document.tables[0].cell(7,3).paragraphs[0].add_run('不正常')
			run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
		else:
			document.add_paragraph('PGA无溢出，表现正常',style="List Bullet 2")
			document.tables[0].cell(7,3).text='正常'
	except:
		document.add_paragraph('No Data found',style="List Bullet 2")
		run=document.tables[0].cell(7,3).paragraphs[0].add_run('No Data found')
		run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

	###################################  内存配置建议 ##########################################
	document.add_heading('4.内存配置建议',level=2)
	GetGraph(args[0],'MTTR advisory',['Size for Est (M)','Estimated Phys Reads (thousands)'],'Buffer Pool Advisory',('buffer cache',DEFAULT_BUFFER_CACHE))
	GetGraph(args[0],'shared pool advisory',['Shared Pool Size(M)','Est LC Load Time (s)'],'Shared Pool Advisory',('shared pool',EST_LC_SIZE))
	GetGraph(args[0],'PGA memory advisory',['PGA Target Est (MB)','Estd PGA Overalloc Count'],'PGA memory advisory',('PGA',PGA_USE))
	GetGraph(args[0],'SGA target advisory for different SGA target sizes',['SGA Target Size (M)','Est Physical Reads'],'SGA Target Advisory',('SGA',SGA_VALUE))
	if advise_flag:
		document.add_paragraph("建议:",style="Heading 9")
		document.add_paragraph("观察内存建议曲线变化趋势，找出拐点，衡量突变点投入产出比，比如增加少量buffer cache值，可以大幅减少物理读，则建议调整，如果增加大量buffer cache值，只能略微减少物理读，则不建议调整。",style="List Bullet 2")
	else:
		pass
# 第五部分，会话登陆阶段
def TimeModelStat(*args):
	# 获取connection management call elapsed time的times和% of DB Time值(TIMES/DB_TIME_TIMEMODEL)
	TIMES=DB_TIME_TIMEMODEL=LOGONS_PER_SEC=LOGONS_PER_TRAN='N/A'
	try:
		*_,tmp=functionC(args[0],"time model statistics")
		for i in tmp:
			if "connection management call elapsed time" in i:
				TIMES=i[1].strip()
				DB_TIME_TIMEMODEL=i[2].strip()
	except:
		pass
	# 获取Load Profile的Logons的Per Second和Per Transaction值
	try:
		*_,tmp=functionC(args[0],"load profile")
		for i in tmp:
			if "Logons:" in i: 
				LOGONS_PER_SEC=i[1].strip()
				LOGONS_PER_TRAN=i[2].strip()
	except:
		pass

	t1=["Statistic Name","Time (s)",r"% of DB Time"]
	t2=["connection management call elapsed time",TIMES,DB_TIME_TIMEMODEL]
	t3=["指标","Per Second","Per Transaction"]
	t4=["Logons",LOGONS_PER_SEC,LOGONS_PER_TRAN]
	# awr报告解析总结 会话登录阶段 会话连接时间的结果
	try:
		TimeModelSug="当前会话连接时间占DB Time比例为%s%%,"%DB_TIME_TIMEMODEL
		if (float(DB_TIME_TIMEMODEL)>1) :
			TimeModelSug+="超过参考值1%，会话连接性能不正常。建议检查数据库连接配置情况，如密码错误，短连接，防火墙，登录访问策略等。"
		else:
			TimeModelSug+="低于参考值1%,会话连接性能正常。"
	except:
		TimeModelSug='No Data found'

	# awr报告解析总结 会话登录阶段 登录次数的结果
	try:
		LogonsSug="当前会话登录每秒为%s个,"%LOGONS_PER_SEC
		if (float(LOGONS_PER_SEC)>80):
			LogonsSug+="超过参考值80个，登录应用连接数量不正常 ，建议尽量减少登录频率，建议使用长连接。"
		else:
			LogonsSug+="少于参考值每秒80个，其每秒登录数在监听处理范围内。"
	except:
		LogonsSug = "No Data found"

	functionB(DB_TIME_TIMEMODEL,1,(8,3),1)
	functionB(LOGONS_PER_SEC,80,(9,3),1)

	document.add_heading("五.会话登录阶段")
	document.add_heading('1.时间模型',level=2)
	generate_table(document,3,t1,[t2],autofit=1)
	document.add_paragraph("建议:",style="Heading 9")
	document.add_paragraph(TimeModelSug,style="List Bullet 2")
	document.add_heading("2.指标",level=2)
	generate_table(document,3,t3,[t4],autofit=1)
	document.add_paragraph("建议:",style="Heading 9")
	document.add_paragraph(LogonsSug,style="List Bullet 2")

# 第六部分，sql解析阶段
def TimeModel(*args):
	# 1.时间模型,每行的值存在timemodelstatlist中
	document.add_heading("六.SQL解析阶段")
	document.add_heading('1.时间模型',level=2)	
	try:
		contextlist=["parse time elapsed","hard parse elapsed time","failed parse elapsed time","hard parse (sharing criteria) elapsed time",\
					"hard parse (bind mismatch) elapsed time"]
		t1=['Statistic Name','Time (s)',r'% of DB Time']

		*_,zjc = functionC(args[0],"time model statistics")
		timemodelstatlist=[i[:3] for i in zjc if i[0] in contextlist]

		Hard_Parse=float(timemodelstatlist[1][1])
		Hard_Fail_Parse=(reduce(lambda x,y:x+y,[float(i[1]) for i in timemodelstatlist if i[0] in ['hard parse elapsed time','failed parse elapsed time']] ))
		Parse_Time=float(timemodelstatlist[0][1])
		Parse_DBtime=float(timemodelstatlist[0][2])
		generate_table(document,3,t1,timemodelstatlist,autofit=1)

		try:
			sugg1="当前失败解析时间与硬解析时间之和占比总解析时间为:%s"%round(100*Hard_Fail_Parse/Parse_Time,3)
			if 100*(Hard_Fail_Parse/Parse_Time)<10 :
				sugg1+="%，小于建议值10%，SQL解析正常。"
			else:
				sugg1+="%，大于建议值10%，SQL解析存在问题，建议检查是否存在异常等待事件，业务sql是否使用绑定变量，是否存在高版本SQL。"
		except:
			sugg1="parse time为0。"
		try:
			sugg2="解析时间占DB Time比例为%s%%，"%Parse_DBtime
			sugg2+="超过参考值20%，解析不正常。" if Parse_DBtime > 20  else "低于参考值20%，解析正常。"
		except:
			sugg2='N/A'
		
		document.add_paragraph("建议:",style="Heading 9")
		document.add_paragraph(sugg1,style="List Bullet 2")
		document.add_paragraph(sugg2,style="List Bullet 2")
	except:
		document.add_paragraph("No Data found",style="List Bullet 2")

	# 2.指标，每行的值存在loadproflist中
	document.add_heading("2.指标",level=2)
	try:
		t2=['指标','Per Second','Per Transaction','ms/次']
		contextlist2=["Parses","Hard parses"]
		######################################   获得总解析次数和总的硬解析次数  ##############################################
		*_,zjc = functionC(args[0],"load profile")
		loadproflist=[i[:3] for i in zjc if i[0].replace('\n','').replace('(SQL):','').replace(':',"").strip() in contextlist2]
		
		*_,zjc = functionC(args[0],"This table displays Key Instance activity statistics")
		parse_count = [i for i in zjc if i[0] in "parse count (total)"]

		*_,zjc = functionC(args[0],"This table displays non-key Instance activity statistics")
		hard_count = [i for i in zjc if i[0] in "parse count (hard)"] 
		######################################   获得总解析次数和总的硬解析次数(当数据库版本为11.2.0.3.0的情况)  ###############
		if parse_count and hard_count:
			pass
		else:
			*_,zjc = functionC(args[0],"This table displays Instance activity statistics")
			parse_count = [i for i in zjc if i[0] in "parse count (total)"]
			hard_count = [i for i in zjc if i[0] in "parse count (hard)"] 


		# 添加第四列的值
		try:
			# print(Parse_Time,parse_count)
			loadproflist[0].append(round((1000*Parse_Time)/float(parse_count[0][1]),3))
		except:
			loadproflist[0].append('N/A')
		try:
			# print(Hard_Parse,hard_count)
			loadproflist[1].append(round((1000*Hard_Parse)/float(hard_count[0][1]),3))
		except:
			loadproflist[1].append("N/A")


		# sql解析阶段  解析的值
		try:
			sugg2="当前解析时间为：%s ms/次%s"%(loadproflist[0][3],("，低于参考值2ms/次，解析正常。" if loadproflist[0][3] < 2 else "，高于参考值2ms/次，解析偏慢。"))
		except:
			sugg2="当前解析数据有问题"

		#sql解析阶段  硬解析的值
		try:
			sugg3="当前硬解析时间为:%s ms/次%s"%(loadproflist[1][3],("，低于参考值5ms/次，硬解析正常。" if loadproflist[1][3] < 5 else "，高于参考值5ms/次，硬解析偏慢。"))
		except:
			sugg3="当前硬解析数据有问题"

		functionB(loadproflist[0][3],2,(10,3),1)
		functionB(loadproflist[1][3],5,(11,3),1)
		generate_table(document,4,t2,loadproflist,autofit=1)
		document.add_paragraph("建议:",style="Heading 9")
		document.add_paragraph(sugg2,style="List Bullet 2")
		document.add_paragraph(sugg3,style="List Bullet 2")
	except:
		document.add_paragraph("No Data found",style="List Bullet 2")	

	# 3.异常等待事件
	soup=args[0].find_all("table",summary=re.compile("Foreground Wait Events and their wait statistics"))
	soup1=args[0].find_all("table",summary=re.compile("background wait events statistics"))
	document.add_heading("3.主要等待事件",level=2)
	# fore_wait_events=[]
	t3=["等待事件","Waits","Total Wait Time (sec)","% DB time","Wait Avg(ms)"]
	contextlist3=["library cache load lock","library cache lock","library cache pin","library cache: mutex S","library cache: mutex X","row cache lock",\
					"cursor: mutex S","cursor: mutex X","cursor: pin S wait on X","cursor: pin S","cursor: pin X","latch: row cache objects","latch: shared pool"]
	tmp2=[]
	tmp1=[]
	try:
		# Foreground Wait Events
		for i in soup[0].find_all('td'):
			if i.string.strip().replace('\n','') in contextlist3: 
				wait_event=i.string.replace('\n','')
				wait_waits=float(i.next_sibling.string.strip().replace(',',''))
				wait_totaltime=float(i.next_sibling.next_sibling.next_sibling.string.strip().replace(',',''))
				wait_dbtime=float(i.next_sibling.next_sibling.next_sibling.next_sibling.next_sibling.next_sibling.string.strip().replace(',',''))
				tmp2.append([wait_event,wait_waits,wait_totaltime,wait_dbtime])
			else:
				pass
		#  background wait events
		for j in soup1[0].find_all('td'):
			if j.string.strip().replace('\n','') in contextlist3: 
				wait_event=j.string.replace('\n','')
				wait_waits=float(j.next_sibling.string.strip().replace(',',''))
				wait_totaltime=float(j.next_sibling.next_sibling.next_sibling.string.strip().replace(',',''))
				try:
					wait_dbtime=float(i.next_sibling.next_sibling.next_sibling.next_sibling.next_sibling.next_sibling.string.strip().replace(',',''))
				except:
					pass
				tmp1.append([wait_event,wait_waits,wait_totaltime,wait_dbtime])
			else:
				pass

		# tmp1中若存在和tmp2，则和tmp2中的值进行相加，否则自己获得平均相应时间之后加到tmp2中
		for i in tmp1:
			k=1
			try:
				for j in tmp2:
					if i[0] in j:
						j[1]+=i[1]
						j[2]+=i[2]
						K=0
					else:
						pass
				if k:
					tmp2.append(i)
			except:
				pass
		# 进行排序和筛选，取total wait time前五的值
		tmp=[]
		fore_wait_events=(sorted(tmp2,key=lambda x:x[2],reverse=True)[:5])
		for i in fore_wait_events:
			i.append(round(1000*float(i[2])/float(i[1]),3))
			# tmp.append(round(1000*float(i[2])/float(i[1]),3))

		hard_parse_list=['latch: row cache objects','latch: shared pool','row cache lock','cursor: mutex X','library cache: mutex X','cursor: pin X','library cache lock']
		soft_parse_list=['library cache: mutex S','cursor: pin S','cursor: mutex S']
		soft_soft_list=['cursor: pin S']
		
		generate_table(document,5,t3,fore_wait_events,autofit=1)
		document.add_paragraph("建议:",style="Heading 9")
		flag=0
		for i in fore_wait_events:
			if i[4] > 1:
				sugg="当前%s单次响应时间为:%sms，超过参考值1ms。"%(i[0],i[4])
				if "latch: row cache objects" in i[0]:
					sugg+="该等待事件一般是由于SQL未使用绑定变量，多版本，执行计划异常导致。当这个等待事件大规模出现时，通常意味着当前数据库性能急剧下降。"
				elif "latch: shared pool" in i[0]:
					sugg+="该等待事件一般是由于shared pool不足，碎片严重，SQL硬解析导致(如未使用绑定变量，多版本等原因)。"
				elif "row cache lock" in i[0]:
					sugg+="该等待事件一般是由于sequence频繁访问，sequence cache过小导致。"
				elif "cursor: mutex X" in i[0]:
					sugg+="该等待事件一般出现在硬解析阶段，常由于sql未使用绑定变量，sql高版本，高频DDL导致。"
				elif "library cache: mutex X" in i[0]:
					sugg+="该等待事件一般出现在硬解析阶段，常由于sql未使用绑定变量，sql高版本，高频DDL导致。"
				elif "cursor: pin X" in i[0]:
					sugg+="该等待事件一般出现在硬解析阶段，常由于sql未使用绑定变量，sql高版本，高频DDL导致。"
				elif "library cache lock" in i[0]:
					sugg+="该等待事件一般出现在硬解析阶段，常由于sql未使用绑定变量，sql高版本，高频DDL导致。"
				elif "library cache: mutex S" in i[0]:
					sugg+="该等待事件一般由于sql版本，长sql，巨量绑定变量导致。"
				elif "cursor: pin S" in i[0]:
					sugg+="该等待事件一般出现在软解析阶段，常由于并发量增大导致。"
				elif "cursor: mutex S" in i[0]:
					sugg+="该等待事件一般出现在软解析阶段，常由于并发量增大导致。"
				document.add_paragraph(sugg,style="List Bullet 2")	
				flag+=1			
			else:
				pass
		if flag == 0:
			document.add_paragraph("无",style="List Bullet 2")			
	except:
		document.add_paragraph("No Data found",style="List Bullet 2")

	# 4.解析数最高的SQL
	document.add_heading("4.解析数最高的SQL",level=2)
	try:
		*a,tmp=functionC(args[0],"top SQL by number of parse calls")
		SQL_TEXT=sorted(tmp,key=lambda x:float(x[0]),reverse=True)[:5]
		SQL_ParseCall = [i[:4] for i in SQL_TEXT]

		cursor_cache=[]
		for i in SQL_ParseCall[:5]:
			try:
				cursor_cache.append(round(100*float(i[0])/float(i[1]),3))
			except:
				cursor_cache.append(0)

		t5=["Parse Calls","Executions","% Total Parses","SQL Id"]
		generate_table(document,4,t5,SQL_ParseCall)
		document.add_paragraph("建议:",style="Heading 9")
		document.add_paragraph("TOP SQL的解析调用为:%s次，SQL ID为%s，SQL Text为%s。"%(SQL_ParseCall[0][0],SQL_ParseCall[0][3],SQL_TEXT[0][-1]),style="List Bullet 2")
		if max(cursor_cache) > 50:
			document.add_paragraph("检查高解析SQL，若存在Parse Calls/Executions大于50%，则说明未良好使用游标缓存功能(游标缓存：oracle建议在中间层延缓关闭常用游标时间，此游标再次被执行时，不需要解析阶段，可直接绑定执行)。",style="List Bullet 2")
		else:
			pass
	except:
		document.add_paragraph("No Data found",style="List Bullet 2")	

	# 5.版本数最高的SQL
	document.add_heading("5.版本数最高的SQL",level=2)
	try:
		*a,tmp=functionC(args[0],"top SQL by version counts",th_len=3)
		SQL_VersionCount=sorted(tmp,key=lambda x:float(x[0]),reverse=True)[:5]

		t5=["Version count","Executions","SQL Id"]
		generate_table(document,3,t5,SQL_VersionCount)
		document.add_paragraph("建议:",style="Heading 9")
		document.add_paragraph("高版本会导致SQL解析效率下降，影响性能，通常是由于绑定变量MISMATCH或者oracle 11G ACS自适应游标共享新特性导致，具体原因可以从V$SQL_SHARED_CURSOR视图中查看高版本原因。",style="List Bullet 2")
		document.add_paragraph("针对绑定变量MISMATCH，一般可以在应用端调大绑定变量所占内存空间解决。针对ACS，一般可以直接关闭新特性解决。关闭方法:",style="List Bullet 2")
		document.add_paragraph('	☆  .alter system set "_optimizer_adaptive_cursor_sharing"=false scope=both;')
		document.add_paragraph('	☆  .alter system set "_optimizer_extended_cursor_sharing_rel"=none  scope=both;')
	except:
		document.add_paragraph("No Data found",style="List Bullet 2")	

# 七.SQL执行阶段
def SqlExecuteTime(*args):
	# 1.第一部分，时间模型
	document.add_heading("七.SQL执行阶段")
	document.add_heading('1.时间模型',level=2)
	try:
		t1=["Statistic Name","Time (s)",r"% of DB Time"]
		*_ ,tmp1 = functionC(args[0],"time model statistics")

		for i in tmp1:
			if "sql execute elapsed time" in i:
				SQL_EXECUTE_TIME=i[:3]
			else:
				pass

		generate_table(document,3,t1,[SQL_EXECUTE_TIME],autofit=1)
		document.add_paragraph("建议:",style="Heading 9")
		sug1="当前SQL执行时间占DB Time比例为%s%%，"%SQL_EXECUTE_TIME[2]
		if float(SQL_EXECUTE_TIME[2])>80:
			sug1 += "大于参考值80%，SQL执行正常。"
		else:
			sug1 += "小于参考值80%，SQL执行不正常，总体性能可能会比较差。"

		document.add_paragraph(sug1,style="List Bullet 2")
		functionB(SQL_EXECUTE_TIME[2],80,(12,3),0)
	except:
		document.add_paragraph("No Data found",style="List Bullet 2")
		functionB([],80,(12,3),0)	

	# 2.指标
	t2=['指标','Per Second','Per Transaction']
	document.add_heading("2.指标",level=2)
	SQL_EXEC_LIST=[]
	contextlist2=["Executes","Logical read","Physical read","Physical write","Read IO","Write IO","Logical reads","Physical reads","Physical writes"]
	*_ ,tmp02 = functionC(args[0],"load profile")

	for i in tmp02:
		if i[0].replace('(SQL)','').replace(':',"").replace('(blocks)','').strip() in contextlist2:
			SQL_EXEC_LIST.append(i[:3])
		elif i[0] in "DB CPU(s):":
			DB_Time = CovertUsToMs(i[1])
		else:
			pass

	t=sorted(SQL_EXEC_LIST)

	generate_table(document,3,t2,SQL_EXEC_LIST,autofit=1)
	document.add_paragraph("建议:",style="Heading 9")
	try:
		tmp=round(1000*DB_Time/float(t[1][1]),3)
		sug1="当前每次逻辑读的响应时间为%sms，%s"%(tmp,("超过建议值0.01ms，逻辑读不正常。" if tmp > 0.01 else "低于建议值0.01ms，逻辑读正常。"))
		functionB(tmp,0.01,(13,3),1)
	except:
		sug1 = "逻辑读数据存在问题"
		functionB([],5,(13,3),1)

	try:
		tmp=round(1000*DB_Time/float(t[2][1]),3)
		sug2="当前每次物理读的响应时间为%sms，%s"%(tmp,("超过建议值10ms，物理读不正常。" if tmp > 10 else "低于建议值10ms，物理读正常。"))
		functionB(tmp,10,(14,3),1)
	except:
		sug2 = "物理读正常。"
		functionB(8,10,(14,3),1)		

	document.add_paragraph(sug1,style="List Bullet 2")
	document.add_paragraph(sug2,style="List Bullet 2")

	# 3.异常等待事件
	document.add_heading("3.主要等待事件",level=2)
	try:
		t3=["等待事件","Waits","Total Wait Time (sec)","Wait Avg(ms)","% DB time"]
		contextlist3=["latch: cache buffers chains","latch: checkpoint queue latch","buffer busy waits","read by other session","db file sequential read","db file scattered read"]

		*_ ,zjc = functionC(args[0],"Foreground Wait Events")
		zjctest=[[i[0],float(i[1]),float(i[3]),CovertUsToMs(i[4]),float(i[-1])] for i in zjc if i[0] in contextlist3]

		fore_wait_events=(sorted(zjctest,key=lambda x:x[2],reverse=True)[:5])
		generate_table(document,5,t3,fore_wait_events,autofit=1)
		document.add_paragraph("建议:",style="Heading 9")
		for i in fore_wait_events:
			if "latch: cache buffers chains" in i:
				document.add_paragraph("%s:Oracle热点块等待事件，由于大量的并发访问同一块造成。一般由于逻辑读过大导致，通常是由索引列不合理导致逻辑读增加，或者执行计划变化导致SQL走全扫描等原因。"\
											%i[0],style="List Bullet 2")
			elif "buffer busy waits" in i:
				document.add_paragraph("%s:竞争类等待事件，代表Oracle遇到竞争或阻塞。此等待事件表示Oracle遇到热点块竞争。是由于DML和DML或者DML和SELECT并发的操作同一块导致。通常的原因是并发DML和Select增高，或者DML操作速度下降。一般由于执行计划异常变动变更导致SQL效率下降，引发应用累积，造成并发量增高。或是较慢的Redo系统，拖慢了DML速度导致。"\
											%i[0],style="List Bullet 2")
			elif "read by other session" in i:
				document.add_paragraph("%s:I/O相关、属于热点块等待事件，多个会话同时物理读同一块导致。一般是由于执行计划异常变更变更导致SQL执行效率下降导致，或存储性能存在问题。"\
											%i[0],style="List Bullet 2")
			elif "db file sequential read" in i:
				document.add_paragraph("%s:Oracle主要的I/O等待事件，单块读等待，在访问索引或以ROWID访问表时出现。如果此等待事件响应超过20ms，有可能是SQL执行计划可能存在变动，使用了选择性较差的索引扫描，导致物理读增加。或者存储性能可能存在问题。"
											%i[0],style="List Bullet 2")
			elif "db file scattered read" in i:
				document.add_paragraph("%s:主要的I/O等待事件，多块读等待，在全表扫描或者索引快速全扫时出现。如果此等待事件响应时间超过20ms，则存储性能可能存在问题，或者SQL执行计划存在异常变动导致物理读增加。对于SQL问题，可以通过创建合理的索引、优化sql（调整SQL所使用的索引）减少全扫操作、解决此类等待事件。"
											%i[0],style="List Bullet 2")
			elif "latch: redo copy" in i:     #当等待时间超过5ms则认为不正常
				sug="latch: redo copy:Log Buffer相关Latch的等待事件，多个Oracle进程同时向Log Buffer写Redo造成竞争时的等待。一般是由于日志量大、并发操作Log Buffer的会话过多导致。当前latch: redo copy的平均响应时间为%sms，"%i[3]
				if i[3] > 5:
					sug+="高于参考值5ms，Log Buffer并发写操作不正常。" 
				else:
					sug+="低于参考值5ms，Log Buffer并发写操作正常。" 
				document.add_paragraph(sug,style="List Bullet 2")
			elif "latch: redo writing" in i:
				sug="Log Buffer相关Latch的等待事件，多个Oracle进程同时呼叫LGWR进程刷新Log Buffer时的等待。一般是由于并发操作Log Buffer的会话过多导致。当前latch: redo writing的平均响应时间为%sms，"%i[3]
				if i[3] > 5:
					sug+="高于参考值5ms，Log Buffer并发写操作不正常。" 
				else:
					sug+="低于参考值5ms，Log Buffer并发写操作正常。" 
				document.add_paragraph(sug,style="List Bullet 2")
			else:
				pass			
	except:		
		document.add_paragraph("No Data found",style="List Bullet 2")	

	# 4.执行时间最长的SQL
	document.add_heading("4.执行事件最长的SQL",level=2)
	try:
		*_ ,zjc = functionC(args[0],"top SQL by elapsed time")
		SQL_ELAPSED_TIME = [i[:7] for i in zjc]
		SQL_TEXT = zjc[0][-1]

		t4=["Elapsed Time (s)","Executions","Elapsed Time per Exec (s)","%Total","%CPU",r"%IO","SQL Id"]
		generate_table(document,7,t4,SQL_ELAPSED_TIME[:5],autofit=1)
		document.add_paragraph("建议:",style="Heading 9")
		document.add_paragraph("TOP SQL的执行时间为:%ss，TOP SQL ID为:%s，TOP SQL TEXT为:%s。"\
								%(SQL_ELAPSED_TIME[0][0],SQL_ELAPSED_TIME[0][-1],SQL_TEXT),style="List Bullet 2")	
		document.add_paragraph("检查总执行时间高及单次执行时间高的SQL，从索引机制，业务逻辑层面优化SQL，降低执行时间。",style="List Bullet 2")	
	except:
		document.add_paragraph("No Data found",style="List Bullet 2")	

	# 5.执行次数最多的SQL
	document.add_heading("5.执行次数最多的SQL",level=2)
	try:
		*_ ,zjc = functionC(args[0],"top SQL by number of executions")
		SQL_Executions = [i[:7] for i in zjc]
		SQL_TEXT = zjc[0][-1]

		t5=["Executions","Rows Processed","Rows per Exec","Elapsed Time (s)","%CPU",r"%IO","SQL Id"]
		generate_table(document,7,t5,SQL_Executions[:5],autofit=1)
		document.add_paragraph("建议:",style="Heading 9")
		document.add_paragraph("TOP SQL总执行次数为%s次，执行频率过高容易引起热点块争用，SQL ID为:%s，SQL TEXT为:%s。"%(\
								SQL_Executions[0][0],SQL_Executions[0][-1],SQL_TEXT),style="List Bullet 2")				
		document.add_paragraph("检查执行次数高的SQL，优化SQL业务逻辑。",style="List Bullet 2")	
	except:
		document.add_paragraph("No Data found",style="List Bullet 2")

	# 6.消耗CPU时间最多的SQL
	document.add_heading("6.消耗CPU时间最多的SQL",level=2)
	try:
		*_ ,zjc = functionC(args[0],"top SQL by CPU time")
		SQL_CPUtime = [i[:4]+i[5:8] for i in zjc]
		SQL_TEXT = zjc[0][-1]

		t6=["CPU Time (s)","Executions","Rows Processed","Rows per Exec","%CPU",r"%IO","SQL Id"]
		generate_table(document,7,t6,SQL_CPUtime[:5],autofit=1)
		document.add_paragraph("建议:",style="Heading 9")
		document.add_paragraph("TOP SQL的CPU消耗时间为:%ss，TOP SQL ID为:%s，TOP SQL TEXT为:%s。"\
								%(SQL_CPUtime[0][0],SQL_CPUtime[0][-1],SQL_TEXT),style="List Bullet 2")	
		document.add_paragraph("检查CPU消耗高的SQL，优化SQL，减少单次执行消耗CPU。建议检查索引机制是否合理，统计信息是否准确，SQL业务逻辑是否可优化。",style="List Bullet 2")
	except:
		document.add_paragraph("No Data found",style="List Bullet 2")	

	# 7.消耗逻辑读最多的SQL
	document.add_heading("7.消耗逻辑读最多的SQL",level=2)
	try:
		*_ ,zjc = functionC(args[0],"top SQL by buffer gets")
		SQL_LOG = [i[:4]+i[5:8] for i in zjc]
		SQL_TEXT = zjc[0][-1]

		t7=["Buffer Gets","Executions","Gets per Exec","%Total","%CPU",r"%IO","SQL Id"]
		generate_table(document,7,t7,SQL_LOG[:5],autofit=1)
		document.add_paragraph("建议:",style="Heading 9")
		document.add_paragraph("TOP SQL的逻辑读为:%s次，TOP SQL ID为:%s，TOP SQL TEXT为:%s。"\
								%(SQL_LOG[0][0],SQL_LOG[0][-1],SQL_TEXT),style="List Bullet 2")
		document.add_paragraph("检查逻辑读高的SQL，优化SQL，减少单次执行消耗逻辑读。逻辑读高的SQL，通常意味着执行计划不合理，也会消耗更多的CPU，建议检查索引机制是否合理，统计信息是否准确，SQL业务逻辑是否可优化。",style="List Bullet 2")	
	except:
		document.add_paragraph("No Data found",style="List Bullet 2")	

	# 8.消耗物理读最多的SQL
	document.add_heading("8.消耗物理读最多的SQL",level=2)
	try:
		*_ ,zjc = functionC(args[0],"top SQL by physical reads")
		SQL_PHY = [i[:4]+i[5:8] for i in zjc]
		SQL_TEXT = zjc[0][-1]

		t8=["Physical Reads","Executions","Reads per Exec","%Total","%CPU",r"%IO","SQL Id"]
		generate_table(document,7,t8,SQL_PHY[:5])
		document.add_paragraph("建议:",style="Heading 9")
		document.add_paragraph("TOP SQL的物理读为:%s次，TOP SQL ID为:%s，TOP SQL TEXT为:%s。"%(\
							SQL_PHY[0][0],SQL_PHY[0][-1],SQL_TEXT),style="List Bullet 2")
		document.add_paragraph("检查物理读高的SQL，优化SQL，减少单次执行消耗物理读。物理读高往往对应大量的全表扫描，建议检查索引机制是否合理，统计信息是否准确，SQL业务逻辑是否可优化。对于频度不高的高物理读消耗SQL，建议选择业务空闲时间段执行。对于不可避免的高物理读消耗SQL，可以通过将表格KEEP到内存中来减少物理读。",style="List Bullet 2")	
	except:
		document.add_paragraph("No Data found",style="List Bullet 2")	

	# 9.逻辑读最多的对象
	document.add_heading("9.逻辑读最多的对象",level=2)
	try:
		*_ ,SQL_LOGICAL_RS = functionC(args[0],"top segments by logical reads",th_len=7)

		t9=["Owner","Tablespace Name","Object Name","Subobject Name","Obj. Type",r"Logical Reads","%Total"]
		generate_table(document,7,t9,SQL_LOGICAL_RS[:5])
		document.add_paragraph("建议:",style="Heading 9")
		document.add_paragraph("逻辑读最多的对象为:%s.%s表空间名为:%s，对象类型为:%s，逻辑读为%s次。"\
								%(SQL_LOGICAL_RS[0][0],SQL_LOGICAL_RS[0][2],SQL_LOGICAL_RS[0][1],SQL_LOGICAL_RS[0][4],SQL_LOGICAL_RS[0][5]),style="List Bullet 2")
		document.add_paragraph("检查逻辑读高的对象，优化对应SQL及应用，降低其逻辑读数量，可以结合逻辑读高的SQL进行分析。",style="List Bullet 2")	
	except:
		document.add_paragraph("No Data found",style="List Bullet 2")	

	# 10.物理读最多的对象
	document.add_heading("10.物理读最多的对象",level=2)
	try:
		*_ ,SQL_physical_RS = functionC(args[0],"top segments by physical reads",th_len=7)

		t10=["Owner","Tablespace Name","Object Name","Subobject Name","Obj. Type","Physical Reads","%Total"]
		generate_table(document,7,t10,SQL_physical_RS[:5])
		document.add_paragraph("建议:",style="Heading 9")
		document.add_paragraph("物理读最多的对象为:%s，表空间名为:%s，对象名为:%s，对象类型为:%s，物理读为%s次。"%(\
								SQL_physical_RS[0][0],SQL_physical_RS[0][1],SQL_physical_RS[0][2],SQL_physical_RS[0][4],SQL_physical_RS[0][5]),style="List Bullet 2")
		document.add_paragraph("检查物理读高的对象，优化SQL及业务，降低或避免物理读，可以结合物理读高的SQL进行分析。如果操作系统有多余的内存及比较空闲的cpu资源，则可以根据实际情况将其keep到内存中。",style="List Bullet 2")	
	except:
		document.add_paragraph("No Data found",style="List Bullet 2")	

	# 11.直接路径读最多的对象
	document.add_heading("11.直接路径读最多的对象",level=2)
	try:
		*_ ,SQL_Dphysical_RS = functionC(args[0],"top segments by direct physical reads",th_len=7)

		t11=["Owner","Tablespace Name","Object Name","Subobject Name","Obj. Type","Direct Reads","%Total"]
		generate_table(document,7,t11,SQL_Dphysical_RS[:5])
		document.add_paragraph("建议:",style="Heading 9")
		document.add_paragraph("直接路径读最多的对象为:%s.%s表空间名为:%s，对象类型为:%s，直接路径读读为%s次。"%(\
							SQL_Dphysical_RS[0][0],SQL_Dphysical_RS[0][2],SQL_Dphysical_RS[0][1],SQL_Dphysical_RS[0][4],SQL_Dphysical_RS[0][5]),style="List Bullet 2")
		document.add_paragraph("检查直接路径读高的对象，直接路径路往往对应并行查询，全表扫描操作，优化相关SQL及应用，降低资源使用率，判断直接路径读参数是否关闭，关闭直接路径读的方法:",style="List Bullet 2")	
		document.add_paragraph('	☆  .alter system set "_serial_direct_read"=never scope=both;')
	except:
		document.add_paragraph("No Data found",style="List Bullet 2")	

# 第八部分:事务提交阶段
def TransCommit(*args):
	# 1.第一部分，时间模型
	# soup=args[0].find_all("table",summary=re.compile("wait class statistics ordered by total wait time"))
	document.add_heading("八.事务提交阶段")
	document.add_heading('1.时间模型',level=2)
	try:
		*_,tmp=functionC(args[0],"foreground wait class statistics")
		for i in tmp:
			if "Commit" in i:
				Tran_Commit=i[:2]+i[3:]
				for i,j in enumerate(Tran_Commit[1:],1):
					try:
						Tran_Commit[i]=CovertUsToMs(j)
					except:
						pass
			else:
				pass

		t1=["Wait Class","Waits","Total Wait Time (sec)","Avg Wait (ms)",r"% of DB Time"]
		sugg1="当前事务提交响应时间为:%sms/次，"%Tran_Commit[3]
		if Tran_Commit[3] > 10 :
			sugg1 += "提交响应时间操作超过参考值10.0ms/次，提交响应时间不正常，提交问题一般由于IO响应慢，日志成员过小，组数过少，日志量异常增大，并发异常增大导致。" 
		else :
			sugg1 += "提交响应时间操作低于参考值10.0ms/次，提交响应时间正常。"

		functionB((Tran_Commit[3]),10,(15,3),1)
		generate_table(document,5,t1,[Tran_Commit])
		document.add_paragraph("建议:",style="Heading 9")
		document.add_paragraph("%s"%sugg1,style="List Bullet 2")
	except:
		document.add_paragraph("No Data found",style="List Bullet 2")
		functionB([],10,(15,3),1)

	# 2.指标,用于获取Redo size (bytes)和Transactions值
	document.add_heading("2.指标",level=2)
	try:
		t2=['指标','Per Second','Per Transaction']
		*_,t=functionC(args[0],"load profile",th_len=3)
		LordProfile=[]
		for i in t:
			if "Redo size" in i[0]:
				LordProfile.append(i)
			elif "Transactions" in i[0]:
				LordProfile.append(i)
				tran_per=i[1]
			elif "DB CPU(s):" in i:
				DB_CPU=i[1]
			else:
				pass
		*_,zjc=functionC(args[0],"Key Instance activity statistics")	
		LordProfile +=[ [i[0],i[2],i[3]] for i in zjc if i[0] in ["user commits","user rollbacks"]]
		
		generate_table(document,3,t2,LordProfile)
		try:
			per_tran_time=round(float(DB_CPU)/float(tran_per),3)
		except:
			per_tran_time='N/A'

		document.add_paragraph("建议:",style="Heading 9")
		document.add_paragraph("关注回滚次数，如果回滚比例过高，说明业务逻辑不合理。",style="List Bullet 2")
		document.add_paragraph("平均事务响应时间为%ss。"%per_tran_time,style="List Bullet 2")
	except:
		document.add_paragraph("No Data found",style="List Bullet 2")	

	# 3.异常等待事件
	document.add_heading("3.主要等待事件",level=2)
	try:
		t3=["等待事件","Waits","Total Wait Time (sec)","Wait Avg(ms)",R"%of DB Time"]
		contextlist3=["latch: redo allocation","latch: redo copy","latch: redo writing","log file sync"]

		*_,zjc=functionC(args[0],"Foreground Wait Events and their wait statistics")
		fore_wait_events=[ [i[0],float(i[1]),float(i[3]),CovertUsToMs(i[4]),i[6]] for i in zjc if i[0] in contextlist3]
		*_,zjc=functionC(args[0],"background wait events statistics")
		fore_wait_events += [ [i[0],float(i[1]),float(i[3]),CovertUsToMs(i[4]),i[6]] for i in zjc if "log file parallel write" in i[0]]

		generate_table(document,5,t3,fore_wait_events)
		document.add_paragraph("建议:",style="Heading 9")
		for i in fore_wait_events:
			if "log file sync" in i[0]:
				sugg="log file sync：提交响应时间等待。等待时间即为事务提交操作的完成时间，通常和log file parallel write一起观察、判断问题。如果log file sync响应时间长而log file"
				sugg += " parallel write不长，通常是并发、Redo切换或log buffer缓存方面存在问题。如果log file sync和log file parallel write响应时间都长，说明是I/O问题，存储性能差导致。"
				sugg += "当前log file sync的平均响应时间为%sms/次"%i[3]
				if i[3] > 10:
					sugg += "高于参考值10ms/次,提交响应时间不正常"
				else:
					sugg += "低于参考值10ms/次,提交响应时间正常"
			elif "latch: redo allocation" in i[0]:
				sugg="latch: redo allocation：Log Buffer相关Latch的等待事件，表示在操作Log Buffer时遇到竞争。一般是由于日志量大、日志写慢，或并发操作Log Buffer的会话过多导致。"
				sugg+="当前latch: redo allocation的平均响应时间为%sms/次"%i[3]
				if i[3] > 10 :
					sugg += "高于参考值5ms/次,日志写不正常"
				else:
					sugg += "低于参考值5ms/次,日志写正常"
			elif "log file parallel write" in i[0]:
				sugg="log file parallel write：Redo File I/O等待事件。一般是由Redo量大或存储性能问题导致。"
				sugg+="当前log file parallel write的平均响应时间为%sms/次"%i[3]
				if i[3] > 5:
					sugg += "高于参考值5ms/次,日志写不正常"
				else:
					sugg += "低于参考值5ms/次,日志写正常"
			else:
				pass
			document.add_paragraph(sugg,style="List Bullet 2")
	except:
		document.add_paragraph("No Data found",style="List Bullet 2")	

	# 4.行级锁最多的对象
	document.add_heading("4.行级锁最多的对象",level=2)
	try:
		*_ ,ROW_LOCK_WAITS = functionC(args[0],"top segments by row lock waits",th_len=7)
		t4=["Owner","Tablespace Name","Object Name","Subobject Name","Obj. Type","Row Lock Waits",r"% of Capture"]
		generate_table(document,7,t4,ROW_LOCK_WAITS[:5])
		document.add_paragraph("建议:",style="Heading 9")
		document.add_paragraph("行级锁最多的对象为:%s.%s，表空间名为:%s，对象类型为:%s，行级锁的数量为%s。"%(\
							ROW_LOCK_WAITS[0][0],ROW_LOCK_WAITS[0][2],ROW_LOCK_WAITS[0][1],ROW_LOCK_WAITS[0][4],ROW_LOCK_WAITS[0][5]),style="List Bullet 2")
		document.add_paragraph("检查行锁最多的对象及对应SQL，检查SQL效率，并发量及业务逻辑，减少单次操作的锁定时间，加快事务提交时间。",style="List Bullet 2")	
	except:
		document.add_paragraph("No Data found")	

# 第九部分RAC Statistics
def RacTotalWaitTime(*args):
	# 1.第一部分，时间模型
	document.add_heading("九.RAC Statistics")
	document.add_heading('1.时间模型',level=2)	
	try:
		# print(rac)
		if rac == 'RAC':
			*_,tmp=functionC(args[0],"foreground wait class statistics")
			for i in tmp:
				if "Cluster" in i:
					RAC_Cluster=i[:2]+i[3:]
				else:
					pass
			t1=["Wait Class","Waits","Total Wait Time (sec)","Avg Wait (ms)",r"%of DB Time"]

			sugg1="当前平均等待时间为:%sms/次，"%RAC_Cluster[3]
			if  CovertUsToMs(RAC_Cluster[3]) > 5:
				sugg1 += "超过参考值5ms/次，集群间响应存在问题，检查集群内部通信网络问题或应用数据访问节点间交互频度。"
			else :
				sugg1 += "低于参考值5ms/次，集群内部通信网络正常。"

			functionB(CovertUsToMs(RAC_Cluster[3]),5,(16,3),5)
			generate_table(document,5,t1,[RAC_Cluster])
			document.add_paragraph("建议:",style="Heading 9")
			document.add_paragraph("%s"%sugg1,style="List Bullet 2")
		else:
			raise
	except:
		document.add_paragraph("No Data found",style="List Bullet 2")
		functionB([],5,(16,3),5)

	# 2.指标
	document.add_heading("2.指标",level=2)
	try:
		t2=['指标','Per Second','Per Transaction']
		contextlist2=["Global Cache blocks received:","Global Cache blocks served:","GCS/GES messages received:","DBWR Fusion writes:","Estd Interconnect traffic (KB)"]

		*_,zjc=functionC(args[0],"global cache load")
		LordProfile = [i for i in zjc if i[0] in contextlist2]

		generate_table(document,3,t2,LordProfile)
		document.add_paragraph("建议:",style="Heading 9")
		document.add_paragraph("当前集群内部网络通信流量为%sKB/s。"%sorted(LordProfile)[1][1],style="List Bullet 2")	
	except IndexError:
		document.add_paragraph("No Data found",style="List Bullet 2")

	# 3.Global Cache and Enqueue Services
	document.add_heading("3.Global Cache and Enqueue Services",level=2)
	try:
		contextlist3=["Avg global cache cr block receive time (ms):","Avg global cache current block receive time (ms):","Avg global cache cr block build time (ms):",\
						"Avg global cache cr block send time (ms):","Global cache log flushes for cr blocks served %:","Avg global cache cr block flush time (ms):",\
						"Avg global cache current block pin time (ms):","Avg global cache current block send time (ms):","Global cache log flushes for current blocks served %:",\
						"Avg global cache current block flush time (ms):"]

		*_,zjc=functionC(args[0],"This table displays workload characteristics for global",th_len=2)
		
		t3=[]
		Global_Cache=[]
		for i in zjc:
			if i[0] in contextlist3:
				Global_Cache.append(i)
				t3.extend(i)
			else:
				pass

		t1=[j for i,j in enumerate(t3) if i%2==0]
		t2=[]
		for j in [j for i,j in enumerate(t3) if i%2==1]:
			try:
				t2.append(float(j))
			except:
				t2.append(0.0)
		t4=dict(zip(t1,t2))
		del t4['Global cache log flushes for current blocks served %:']
		del t4['Global cache log flushes for cr blocks served %:']
		# sugg3="no Data found" if not(t4.values()) else ("平均响应时间超过参考值20ms,节点内部通信性能存在问题,交互太频繁。" if max(t4.values())>20 else "平均响应时间低于参考值20ms，节点内部通信性能正常。")

		generate_table(document,2,['Global Cache and Enqueue Services','数值'],Global_Cache)
		document.add_paragraph("建议:",style="Heading 9")

		flag=0
		for i in Global_Cache:
			if '%' not in i[0]:
				try:
					if 'Avg global cache current block pin time (ms):' in i:
						if float(i[1]) > 8:
							sugg3 ="当前%s的值为%sms，超过参考值8ms，不正常。"%(i[0].replace('(ms)','').replace(':',''),i[1])
							flag=1
							document.add_paragraph("%s"%sugg3,style="List Bullet 2")	
						else:
							pass
					else:
						if CovertUsToMs(i[1]) > 10:
							sugg3 ="当前%s的值为%sms，超过参考值10ms，不正常。"%(i[0].replace('(ms)','').replace(':',''),i[1])
							flag=1
							document.add_paragraph("%s"%sugg3,style="List Bullet 2")	
						else:
							pass
				except:
					pass
			else:
				pass
			
		if flag:
			document.add_paragraph("节点内部网络通信性能不正常。",style="List Bullet 2")
			run=document.tables[0].cell(17,3).paragraphs[0].add_run('不正常')
			run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
		else:
			document.add_paragraph("节点内部网络通信性能正常。",style="List Bullet 2")
			document.tables[0].cell(17,3).text = '正常'
	except:
		document.add_paragraph("No Data found",style="List Bullet 2")
		run=document.tables[0].cell(17,3).paragraphs[0].add_run('N/A')
		run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)


	# 4.节点内部通信PING延迟
	document.add_heading("4.节点内部通信PING延迟",level=2)
	try:
		title,value=functionC(args[0],"IC ping latency statistics")
		# print(1,value,len(title))
		generate_table(document,len(title),title,value)
		# print(title,'\n',value)
		document.add_paragraph("建议:",style="Heading 9")
		if float(value[0][2]) < 1 and float(value[0][5]) < 1 and float(value[1][2]) <1 and float(value[1][5]) <1:
			document.add_paragraph("节点内部通信PING延迟低于参考值1ms，正常。",style="List Bullet 2")	
			document.tables[0].cell(18,3).text = '正常'
		else:
			document.add_paragraph("节点内部通信PING延迟超过参考值1ms，不正常。",style="List Bullet 2")	
			run=document.tables[0].cell(18,3).paragraphs[0].add_run('不正常')
			run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
	except:
		document.add_paragraph("No Data found",style="List Bullet 2")
		run=document.tables[0].cell(18,3).paragraphs[0].add_run('N/A')
		run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

	# 5.Dynamic Remastering Stats 
	document.add_heading("5.Dynamic Remastering Stats",level=2)	
	try:
		title,value=functionC(args[0],"Dynamic Remastering Stats. . times are in seconds")

		generate_table(document,len(title),title,value)	
		document.add_paragraph("建议:",style="Heading 9")
		document.add_paragraph("建议关闭drm,关闭方法:",style="List Bullet 2")
		document.add_paragraph("""	☆  .alter system set "_gc_policy_time"=0 scope=spfile sid='*';""",style="List Bullet 3")
	except:
		document.add_paragraph("无",style="List Bullet 2")

	# 6.异常等待事件
	document.add_heading("6.主要等待事件",level=2)
	try:
		t4=["等待事件","Waits","Total Wait Time (sec)","Wait Avg(ms)",'% DB time']
		contextlist3=["gc cr multi block request","gc buffer busy acquire","gc current block busy","gc cr block busy","gcs log flush sync","gc current multi block request"]

		*_,zjc =functionC(args[0],"Foreground Wait Events and their wait statistics")
		fore_wait_events = [ [i[0],float(i[1]),float(i[3]),CovertUsToMs(i[4]),i[6]] for i in zjc if i[0] in contextlist3]

		if fore_wait_events:
			generate_table(document,5,t4,fore_wait_events,autofit=1)
			mm=[i[3] for i in fore_wait_events]
			if max(mm) > 1 :
				sugg4="节点间通信网络平均响应时间为%sms，超过参考值1.0ms，集群节点间网络通信存在问题或者集群节点交互过于频繁，建议检查节点间通信网络，减少节点间交互。"%max(mm)
			else:
				sugg4="节点间通信网络平均响应时间为%sms，低于参考值1.0ms，集群节点间网络通信正常。"%max(mm)

			document.add_paragraph("建议:",style="Heading 9")
			document.add_paragraph("gc相关的等待都是基于节点间通信网络性能。",style="List Bullet 2")	
			document.add_paragraph(sugg4,style="List Bullet 2")	
		else:
			document.add_paragraph("No Data found",style="List Bullet 2")
	except:
		document.add_paragraph("No Data found",style="List Bullet 2")

	# 7.Global Cache Buffer Busy最多的对象
	document.add_heading("7.Global Cache Buffer Busy最多的对象",level=2)
	try:
		*_,GLOBAL_CACHE_BUFFER_BUSY =functionC(args[0],"top segments by row lock waits",th_len=7)

		t5=["Owner","Tablespace Name","Object Name","Subobject Name","Obj. Type","GC Buffer Busy",r"% of Capture"]
		generate_table(document,7,t5,GLOBAL_CACHE_BUFFER_BUSY[:5])
		document.add_paragraph("建议:",style="Heading 9")
		document.add_paragraph("Global Cache Buffer Busy最多的对象为%s.%s，表空间为%s，对象类型为%s，GC Buffer Busy等待次数为%s。"%(GLOBAL_CACHE_BUFFER_BUSY[0][0],\
								GLOBAL_CACHE_BUFFER_BUSY[0][2],GLOBAL_CACHE_BUFFER_BUSY[0][1],GLOBAL_CACHE_BUFFER_BUSY[0][4],GLOBAL_CACHE_BUFFER_BUSY[0][-2]),style="List Bullet 2")
		document.add_paragraph("检查节点间交互对象情况，检查相关SQL及业务，尽量减少节点间对象交互。",style="List Bullet 2")	
	except:
		document.add_paragraph("No Data found",style="List Bullet 2")

#第十部分：数据库参数配置建议	
def InitParaAdvise(*args):
	document.add_heading("十.数据库参数建议")
	try:
		advise_list=[]
		diff_list=[]
		*_,zjc=functionC(args[0],"name and value of init.ora parameters",zipnum=[0,1])
		t3=dict(zip(zjc[0],zjc[1]))
		key_word=init_para_cmplist.keys()
		for i in key_word:
			t=t3.get(i,'N/A')
			advise_list.append([i,t,init_para_cmplist[i][0]])
			if (t.upper() == init_para_cmplist[i][0]):
				pass
			else:
				diff_list.append([i,init_para_cmplist[i]])

		generate_table(document,3,['参数','当前值','建议值'],advise_list)
		document.add_paragraph("建议:",style="Heading 9")
		if diff_list:
			for i in diff_list:
				document.add_paragraph("建议修改%s的值为%s。%s"%(i[0],i[1][0],i[1][1]),style="List Bullet 2")
			document.tables[0].cell(19,3).text = '存在'
		else:
			document.add_paragraph('无',style="List Bullet 2")
			document.tables[0].cell(19,3).text = '不存在'
	except:
		document.add_paragraph("No Data found",style="List Bullet 2")

def run(*args):
	try:
		if not(os.path.exists(args[0])):
			raise PermissionError
		try:
			soup=BeautifulSoup(open(args[0],'r',encoding='utf-8'),'html.parser')
		except :
			soup=BeautifulSoup(open(args[0],'r',encoding='gbk'),'html.parser')
		MainReport()
		function=[BasicSituation,HostResource,MemoryConfig,TimeModelStat,TimeModel,SqlExecuteTime,TransCommit,RacTotalWaitTime,InitParaAdvise]
		try:
			for i in function:
				i(soup)
		except:
			print("该AWR分析工具暂时只支持Oracle数据库版本为11.2.0.2.0以上版本的单机或rac")
		document.save(file)
		
		print("AWR分析结果保存在%s中"%file)
		# print(file.replace('docx','pdf'))
		t = Doc2PDF(file)
		t.change2pdf()

	except PermissionError:
		print(" AWR-0001:参数错误，未能传入正确的参数值。",'\n',\
			  "AWR-0002:生成的docx文件已打开。")

# 单元测试用
class unit(unittest.TestCase):
	def test_run(self):
		testfile=os.path.join(tmploca,'12c_awrrpt_1_2300_2301.html')
		try:
			soup=BeautifulSoup(open(testfile,'r',encoding='utf-8'),'html.parser')
		except :
			soup=BeautifulSoup(open(testfile,'r',encoding='gbk'),'html.parser')
		global file
		MainReport()

		# function=[BasicSituation,HostResource,MemoryConfig,TimeModelStat,TimeModel,SqlExecuteTime,TransCommit,RacTotalWaitTime,InitParaAdvise]
		function=[TimeModel]
		try:
			for i in function:
				i(soup)
		except:
			print("该awr分析工具暂时只支持Oracle数据库版本为11.2.0.1.0以上版本的单机或rac")

def test():
	testfile=os.path.join(tmploca,'awrrpt_1_25045_25053.html')
	try:
		soup=BeautifulSoup(open(testfile,'r',encoding='utf-8'),'html.parser')
	except :
		soup=BeautifulSoup(open(testfile,'r',encoding='gbk'),'html.parser')
	global file
	file=os.path.join(resloca,os.path.basename(testfile).replace('html','docx'))
	MainReport()
	function=[BasicSituation,HostResource,MemoryConfig,TimeModelStat,TimeModel,SqlExecuteTime,TransCommit,RacTotalWaitTime,InitParaAdvise]
	try:
		for i in function:
			i(soup)

		document.save(file)

	# out = file.replace('docx','pdf')
	# print(file,'\n',out)
	# t = Doc2PDF(file)
	# t.change2pdf()
	
	except:
		print("该AWR分析工具暂时只支持Oracle数据库版本为11.2.0.2.0以上版本的单机或rac")

if __name__ == "__main__":
	run(' '.join(sys.argv[1:]))	
	# unittest.main()
	# test()